import os
import re
import logging
from typing import Dict, Any, List, Optional
from pathlib import Path
import docx

logger = logging.getLogger("docx_processor")

class DocxProcessor:
    """
    Industrial DOCX Document Processing Engine.
    Extracts paragraphs, headings, technical specification tables,
    and industrial attributes with document section grounding.
    """

    DOCX_SPEC_PATTERNS = [
        # Identification
        (r'(?:Model(?:\s*No\.?|\s*Number)?|Part\s*(?:No\.?|Number)|Catalog\s*No\.?|Type\s*Designation)\s*[:=\-]\s*([A-Z0-9\-_/\.]{3,30})', 'Model Identifier'),
        (r'(?:Manufacturer|Brand|Make|Vendor)\s*[:=\-]\s*([A-Za-z0-9\s&\.\-]{3,30})', 'Manufacturer'),
        (r'(?:Product\s*Series|Series)\s*[:=\-]\s*([A-Za-z0-9\-_\s]{3,25})', 'Product Series'),

        # Electrical Ratings
        (r'(?:Rated\s+|Nominal\s+|Output\s+)?Power(?:\s*Rating)?\s*[:=\-]\s*([0-9\.]+\s*(?:kW|HP|W|MW|kVA|VA))', 'Rated Power'),
        (r'(?:Rated\s+|Nominal\s+|Operating\s+)?Voltage\s*[:=\-]\s*([0-9\.\/]+\s*(?:V|kV|VAC|VDC|Volts)(?:\s*[\u00b1\+\-0-9%/\s]+)?)', 'Rated Voltage'),
        (r'(?:Rated\s+|Nominal\s+|Full\s*Load\s+)?Current\s*[:=\-]\s*([0-9\.\/]+\s*(?:A|Amps|mA|Amperes))', 'Rated Current'),
        (r'(?:Nominal\s+|Supply\s+)?Frequency\s*[:=\-]\s*([0-9\.]+\s*(?:Hz|kHz))', 'Frequency'),
        (r'(?:Phase|No\.\s*of\s*Phases)\s*[:=\-]\s*([0-9]+\s*(?:Phase|Ph|\~)|Single\s*Phase|Three\s*Phase|3-Phase|1-Phase)', 'Phase Configuration'),
        (r'(?:Power\s*Factor|Cos\s*[\u03c6\u03d5\u00f8]|PF)\s*[:=\-]\s*([0-9\.]+)', 'Power Factor'),

        # Mechanical & Dynamic Specifications
        (r'(?:Rated\s+|Nominal\s+|Synchronous\s+)?Speed\s*[:=\-]\s*([0-9\.]+\s*(?:RPM|rpm|r/min|min\^\-1))', 'Synchronous Speed'),
        (r'(?:Rated\s+|Nominal\s+|Full\s*Load\s+|Breakdown\s+)?Torque\s*[:=\-]\s*([0-9\.]+\s*(?:Nm|N\-m|lb\-ft|kgf\-m))', 'Rated Torque'),
        (r'(?:Frame\s*(?:Size|Type)?|Housing)\s*[:=\-]\s*([0-9]{2,4}[A-Z0-9\-_/]+)', 'Frame Size'),
        (r'(?:Mounting(?:\s*Type|\s*Arrangement)?)\s*[:=\-]\s*([A-Za-z0-9\-\s\(\)\/]{3,30})', 'Mounting Type'),
        (r'(?:Shaft\s*(?:Diameter|Extension)?)\s*[:=\-]\s*([0-9\.]+\s*(?:mm|inch|in|\"))', 'Shaft Diameter'),
        (r'(?:Total\s+|Gross\s+|Net\s+)?Weight\s*[:=\-]\s*([0-9\.]+\s*(?:kg|lbs|g|tonne|ton))', 'Unit Weight'),

        # Protection, Thermal & Environmental
        (r'(?:Protection(?:\s*Class|\s*Degree|\s*Rating)?|Enclosure|IP\s*Rating|Degree\s*of\s*Protection)\s*[:=\-]\s*(IP\s*[0-9]{2}[A-Z]?|NEMA\s*[0-9A-Z]+)', 'Enclosure Protection'),
        (r'(?:Insulation(?:\s*Class|\s*System)?|Thermal\s*Class)\s*[:=\-]\s*(?:Class\s+)?([A-H]|Class\s+[A-H]|F\s*\(155\s*°C\)|H\s*\(180\s*°C\))', 'Insulation Class'),
        (r'(?:Efficiency(?:\s*Class|\s*Rating|\s*Level)?|IE\s*Class)\s*[:=\-]\s*(IE[1-5]|[0-9\.]+\s*%)', 'Full Load Efficiency'),
        (r'(?:Duty\s*(?:Cycle|Type)?|Rating)\s*[:=\-]\s*(S[1-9]|Continuous|Intermittent|S1\s*Continuous)', 'Duty Cycle'),
        (r'(?:Cooling(?:\s*Method|\s*Type)?)\s*[:=\-]\s*([A-Za-z0-9\-\s]{3,20}|IC[0-9]{3})', 'Cooling Method'),
        (r'(?:Ambient\s*(?:Temperature|Temp)?|Operating\s*Temperature)\s*[:=\-]\s*([\-0-9\.\s\+to°degCF]+)', 'Ambient Temperature'),
        (r'(?:Noise(?:\s*Level|\s*Pressure)?)\s*[:=\-]\s*([0-9\.]+\s*(?:dB\(A\)|dB|dBA))', 'Noise Level'),
        (r'(?:Bearing\s*(?:Type|DE|NDE)?)\s*[:=\-]\s*([A-Za-z0-9\-_]{4,20})', 'Bearing Designation'),

        # Standards & Certifications
        (r'(?:Standard|Compliance|Applicable\s*Standard)\s*[:=\-]\s*([A-Za-z0-9\-\s\/\.\(\)]{3,40})', 'Compliance Standard'),
        (r'(?:Certification|Certificates?)\s*[:=\-]\s*([A-Za-z0-9\s,\/\-\+]+)', 'Certifications'),
        (r'(?:Hazardous\s*Area|ATEX\s*Rating|Explosion\s*Proof)\s*[:=\-]\s*([A-Za-z0-9\s\/\-_]{3,30})', 'Hazardous Area Rating')
    ]

    @staticmethod
    def _clean_key(key: str) -> str:
        k = re.sub(r'[^a-zA-Z0-9\s]', ' ', key).strip()
        words = [w.capitalize() for w in k.split() if len(w) > 0]
        return " ".join(words)

    @staticmethod
    def extract_docx_content(file_path: str, filename: str) -> Dict[str, Any]:
        """
        Process a Word document (.docx), extract text, headings, tables, and specifications.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found at path: {file_path}")

        extracted_attributes = {}
        source_citations = []
        text_stream_parts = []

        try:
            doc = docx.Document(file_path)
            
            # 1. Process Paragraphs
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            full_text_corpus = "\n".join(paragraphs)

            if paragraphs:
                text_stream_parts.append("=== DOCUMENT PARAGRAPHS & HEADINGS ===\n" + "\n".join(paragraphs))

            # 2. Process Tables
            tables_count = len(doc.tables)
            if tables_count > 0:
                text_stream_parts.append(f"\n=== DOCUMENT TABLES ({tables_count} Tables Found) ===")
                
                for t_idx, table in enumerate(doc.tables):
                    table_lines = []
                    table_lines.append(f"\n--- Table {t_idx + 1} ---")
                    
                    for row_idx, row in enumerate(table.rows):
                        cells = [c.text.strip() for c in row.cells]
                        table_lines.append(" | ".join(cells))

                        # If 2-column key-value table, extract attribute
                        if len(cells) >= 2 and cells[0] and cells[1]:
                            k_raw = cells[0].rstrip(':=-').strip()
                            v_raw = cells[1].strip()
                            if 2 < len(k_raw) < 35 and 1 < len(v_raw) < 80:
                                if k_raw.lower() not in ['parameter', 'specification', 'attribute', 'item', 'description', 'property']:
                                    cleaned_k = DocxProcessor._clean_key(k_raw)
                                    if cleaned_k and cleaned_k not in extracted_attributes:
                                        extracted_attributes[cleaned_k] = v_raw
                                        source_citations.append({
                                            "page": 1,
                                            "attribute": cleaned_k,
                                            "snippet": f"Table {t_idx + 1}: {k_raw} -> {v_raw}",
                                            "confidence": 0.98
                                        })

                    text_stream_parts.append("\n".join(table_lines))

            # 3. Regex Specification Extraction across full corpus
            for pattern, attr_name in DocxProcessor.DOCX_SPEC_PATTERNS:
                match = re.search(pattern, full_text_corpus, re.IGNORECASE)
                if match and attr_name not in extracted_attributes:
                    val = match.group(1).strip().replace('\n', ' ')
                    extracted_attributes[attr_name] = val
                    source_citations.append({
                        "page": 1,
                        "attribute": attr_name,
                        "snippet": match.group(0).strip().replace('\n', ' '),
                        "confidence": 0.96
                    })

            # 4. Generic Key-Value Line Extraction
            for p in paragraphs:
                if 4 < len(p) < 100:
                    kv_match = re.match(r'^([A-Za-z0-9\s\/\-\.]{3,30})\s*[:=\-]\s+([A-Za-z0-9\.\,\s\/\-\%\u00b1\u00b0\+\(\)]+)$', p)
                    if kv_match:
                        raw_k = kv_match.group(1).strip()
                        raw_v = kv_match.group(2).strip()
                        if not raw_k.lower().startswith(('doc', 'page', 'http', 'www', 'note', 'tel', 'author')):
                            cleaned_k = DocxProcessor._clean_key(raw_k)
                            if cleaned_k and len(cleaned_k) >= 3 and cleaned_k not in extracted_attributes:
                                extracted_attributes[cleaned_k] = raw_v
                                source_citations.append({
                                    "page": 1,
                                    "attribute": cleaned_k,
                                    "snippet": p,
                                    "confidence": 0.93
                                })

        except Exception as docx_err:
            logger.error(f"DOCX extraction error: {docx_err}")
            raise docx_err

        # Summary attributes
        extracted_attributes["File Type"] = "Microsoft Word Document (.docx)"
        extracted_attributes["Total Paragraphs"] = str(len(paragraphs))
        extracted_attributes["Total Tables"] = str(tables_count)

        summary = f"Processed Word document containing {len(paragraphs)} paragraph(s) and {tables_count} table(s). Extracted {len(extracted_attributes)} specifications with layout grounding."

        return {
            "pages_count": max(1, (len(paragraphs) + tables_count * 5) // 12 + 1),
            "extracted_summary": summary,
            "extracted_attributes": extracted_attributes,
            "source_citations": source_citations,
            "extracted_text": "\n\n".join(text_stream_parts)
        }
